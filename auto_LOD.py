"""
Autogenerates a List of Documents in docx format from a folder containing the appropriate documents.

Each document must be named <date in yyyy.mm.dd format> <time in 24h format (optional)> <Name of document>.<file_extension>

The table will have three columns:
1. Tab Number ('S/N') (list index + 1)
2. Date and Time (time for emails only) ('Date')
4. Name ('Description of Document')
"""

# TODO - Regex filter to capture the names of each

import docx
import os
import re
from collections import namedtuple
import argparse
import sys

lod_files = []
LOD_File = namedtuple('LOD_File', 'index_number date time description extension') 
invalid_names = []


# Eg 2020.01.12 23.59 abcd.pdf OR
# Eg 2020.05.22 abcd.docx
file_name_pattern = r'^([0-9]{4}\.[0-9]{2}\.[0-9]{2})( )(([0-2][0-4]\.[0-5][0-9])( ))?(.+)((\.)(.+))$'

parser = argparse.ArgumentParser(description="Create a LOD file.\n"
                                             "Example: python auto_LOD.py -d ./test_files -n 'New Document.docx'")

# This is the folder where all your named files are located. You can set this to an absolute path.
parser.add_argument('-d', '--directory', metavar='file_directory', type=str,
                    help='A file directory containing all the documents to be listed. Pass in an absolute or relative '
                         'path to the directory.',
                    default='./test_files')

# This is the output file (a .docx or .doc). You can also change this to an absolute directory path
parser.add_argument('-n', '--name', metavar='file_name',
                    type=str,
                    help='A file name or file path for the LOD file.',
                    default='New LOD.docx')

args = parser.parse_args()

lod_files_directory = args.directory
word_doc_file_name = args.name

if not os.path.isdir(lod_files_directory):
    sys.exit(f"ERROR: Directory '{lod_files_directory}' does not exist. Check the file path for any errors.")

# If the word doc file name is on a filepath that does not exist, create the intermediate files
word_doc_file_path_components = os.path.split(word_doc_file_name)
if len(word_doc_file_path_components) >= 2:
    if not os.path.isdir(os.path.join(*word_doc_file_path_components[:-1])):
        os.makedirs(os.path.join(*word_doc_file_path_components[:-1]))
        print(f"Created intermediate directories {word_doc_file_name}...")

if not word_doc_file_name.endswith('.docx') and not word_doc_file_name.endswith('.doc'):
    while os.path.isfile(word_doc_file_name + ".docx"):
        # Add a 1 to the end of the file name if the file already exists.
        print(f"File {word_doc_file_name} already exists.")
        word_doc_file_name += "1"
        print(f"Output file renamed to {word_doc_file_name}")

    # If word doc file name does not have .docx at the end, add it.
    word_doc_file_name += ".docx"
    print(f"Adding .docx to Word Document file name. File name changed to {word_doc_file_name}")

else:
    # Add a 1 to the end of the file name if the file already exists.
    while os.path.isfile(word_doc_file_name):
        print(f"{word_doc_file_name} already exists.")
        file_name_extractor_pattern = r'^(.*)\.(doc|docx)$'
        word_doc_file_name_before_extension = re.match(file_name_extractor_pattern, word_doc_file_name)[1]
        word_doc_file_name_extension = re.match(file_name_extractor_pattern, word_doc_file_name)[2]
        word_doc_file_name = word_doc_file_name_before_extension + "1" + "." + word_doc_file_name_extension
        print(f"Output File renamed to {word_doc_file_name}")


def continue_or_abort(prompt="Continue (Y) or Abort (N)?", abort_message="Aborted. Rename Files"):
    """
    Raise an exception if the user does not choose to continue after encountering an invalid file name.
    """
    continue_ = input(prompt)
    if continue_.lower() == 'abort':
        raise Exception(abort_message)
    elif continue_.lower().startswith('y') or continue_.lower() == 'continue':
        return
    else:
        raise Exception(abort_message)


for index, entry in enumerate(os.scandir(lod_files_directory)):
    file_name = entry.name
    match = re.match(file_name_pattern, file_name)

    # If there is no regex name, prompt to rename the file and continue the loop
    if not match:
        invalid_names.append(file_name)
        # If user chooses to break, an error will be raised. Otherwise, continue.
        continue_or_abort(prompt=f"Invalid file name {file_name} (does not follow pattern). Continue (Y) or abort (N)? ",
                          abort_message="Aborted. Rename files.")
        continue
    # Else, add each group to a tuple of format date, time, name
    elif match:
        index_number = index + 1
        # Extract from regex the relevant parameters
        date = match.group(1)
        time = match.group(4)
        description = match.group(6)

        # Saving this just in case
        extension = match.group(9)

        # Save a named tuple to the lod_files list
        lod_files.append(LOD_File(index_number, date, time, description, extension))

# no_of_files = len(lod_files)

# Alert user that the file names are wrongly formatted
if invalid_names:
    print("Rename the following files:")
    for invalid_name in invalid_names:
        print(invalid_name)
    continue_or_abort(prompt=f"Skip these files and continue (Y) or abort (N)? ",
                      abort_message="Aborted. Rename files.")

# Print file names to console
print(f"\nCreating LOD file from {os.path.abspath(lod_files_directory)} at {os.path.abspath(word_doc_file_name)}, "
      f"containing the following files:\n")

for lod_file in lod_files:
    print(f"{lod_file.index_number} | {lod_file.date}{' ' + lod_file.time if lod_file.time else ''} | {lod_file.description}")

# Create a new  Microsoft Word document
lod_doc = docx.Document()

# Create a table with just headers first. Add rows as we iterate through loop.
table = lod_doc.add_table(1, 3)

# Make sure there are borders
table.style = 'Table Grid'

# Populate header row

heading_cells = table.rows[0].cells
heading_cells[0].text = 'S/N'
heading_cells[1].text = 'Date/Time'
heading_cells[2].text = 'Description'

# TODO - change this later
# # Either "yyyy.mm.dd" or "day month year"
# DATE_FORMAT = "day month year"

# # Either "24h" or "12h"
# TIME_FORMAT = "12h"

LIST_OF_MONTHS = [
    "January",
    "February",
    "March",
    "April",
    "May",
    "June",
    "July",
    "August",
    "September",
    "October",
    "November",
    "December"
]

for lod_file in lod_files:
    cells = table.add_row().cells

    # Adds a list number to S/N column
    cells[0].text = ""  # f"{lod_file.index_number}"
    cells[0].paragraphs[0].style = 'List Number'
    yyyymmdd_date = lod_file.date

    components_of_date = yyyymmdd_date.split(".")
    year = int(components_of_date[0])
    month = int(components_of_date[1])
    day = int(components_of_date[2])

    month_in_words = LIST_OF_MONTHS[month - 1]

    word_date = str(day) + " " + month_in_words + " " + str(year)

    if lod_file.time is not None:
        time_24h = lod_file.time
        components_of_time = time_24h.split(".")
        hour = int(components_of_time[0])
        
        # Minute is not changed to int because it would otherwise show things like "10.0 pm"
        minute = components_of_time[1]
        am_or_pm = ""
        if hour < 12:
            am_or_pm = "am"
            if hour != 0:
                hour_12h = hour
            elif hour == 0:
                hour_12h = 12
            
        elif int(hour) >= 12:
            am_or_pm = "pm"
            if hour != 12:
                hour_12h = hour - 12
            elif hour == 12:
                hour_12h = hour
        
        time_12h = f"{hour_12h}.{minute}{am_or_pm}"

    if lod_file.time is not None:
        cells[1].text = f"{word_date}\n{time_12h}"
        
    elif lod_file.time is None:
        cells[1].text = word_date
    cells[2].text = f"{lod_file.description}"

try:
    lod_doc.save(word_doc_file_name)
except Exception as e:
    print(f"{type(e)}: {e}")
    print("Failed to save new LOD.")
else:
    print(f"\nSuccessfully saved new list of documents\n"
          f"Folder Used: '{os.path.abspath(lod_files_directory)}'\n"
          f"Output File Path: '{os.path.abspath(word_doc_file_name)}'")
